"""阶段 1 冒烟测试：QT_QPA_PLATFORM=offscreen 下跑，不弹 GUI。

用法：.venv/Scripts/python.exe scripts/smoke_phase1.py
"""
import os
import sys
import tempfile

# 必须在 import 任何 app 模块前设置
_tmp = tempfile.mkdtemp(prefix="nailong_smoke_")
os.environ["NAILONG_DATA_DIR"] = _tmp
os.environ.setdefault("QT_QPA_PLATFORM", "offscreen")

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

PASS = []


def check(name, cond):
    status = "PASS" if cond else "FAIL"
    PASS.append(cond)
    print(f"[{status}] {name}")
    if not cond:
        raise AssertionError(name)


# ---------- db ----------
from app.db import Database, data_dir
from app.models import Editor, Manuscript, Submission, Reply

db = Database()
check("数据目录覆盖生效", data_dir() == _tmp)
check("files 子目录已创建", os.path.isdir(os.path.join(_tmp, "files")))

eid = db.insert_editor(Editor(name="张三编辑", platform="平台A", email="ed1@x.com",
                              genres="言情/悬疑", fee_info="千字100", notes="测试"))
eid2 = db.insert_editor(Editor(name="李四编辑", platform="平台B", email="ed2@x.com",
                               genres="科幻", favorite=True))
check("editor 插入", eid > 0 and eid2 > 0)
check("editor 列表/筛选", len(db.list_editors(keyword="张三")) == 1
      and len(db.list_editors(platform="平台B")) == 1
      and len(db.list_editors(genre="悬疑")) == 1
      and len(db.list_editors(favorites_only=True)) == 1)
ed = db.get_editor(eid)
ed.fee_info = "千字200"
db.update_editor(ed)
check("editor 更新", db.get_editor(eid).fee_info == "千字200")
check("toggle_favorite", db.toggle_favorite(eid) is True and db.toggle_favorite(eid) is False)
check("toggle_blacklisted", db.toggle_blacklisted(eid) is True
      and all(e.id != eid for e in db.list_editors())
      and len(db.list_editors(include_blacklisted=True)) == 2)
db.toggle_blacklisted(eid)
check("distinct_platforms", db.distinct_platforms() == ["平台A", "平台B"])
check("distinct_genres", "言情" in db.distinct_genres() and "科幻" in db.distinct_genres())

mid = db.insert_manuscript(Manuscript(title="测试小说", file_path="a.docx", word_count=12345,
                                      category="短篇", reader_group="女频", emotion="甜",
                                      style="轻松", genre_type="现言"))
check("manuscript CRUD",
      db.get_manuscript(mid).word_count == 12345 and len(db.list_manuscripts()) == 1)
m = db.get_manuscript(mid)
m.word_count = 13000
db.update_manuscript(m)
check("manuscript 更新", db.get_manuscript(mid).word_count == 13000)

sid = db.insert_submission(Submission(manuscript_id=mid, editor_id=eid,
                                      from_mailbox="me@qq.com", to_email="ed1@x.com",
                                      subject="投稿《测试小说》", body="正文"))
db.update_status(sid, "已发")
check("submission 已发", db.list_submissions()[0].status == "已发"
      and db.list_submissions()[0].sent_at != "")
check("find_pending 一稿一投", db.find_pending(mid, eid) is not None
      and db.find_pending(mid, eid2) is None)
db.update_reply_status(sid, "过稿")
check("update_reply_status 后不再 pending", db.find_pending(mid, eid) is None)
check("count_editor_last_days", db.count_editor_last_days(eid, 7) == 1)
check("count_today", db.count_today("me@qq.com") == 1 and db.count_today("other@qq.com") == 0)
check("status_filter", len(db.list_submissions(status_filter="已发")) == 1
      and len(db.list_submissions(status_filter="失败")) == 0)

rid = db.insert_reply(Reply(submission_id=sid, from_email="ed1@x.com", subject="Re: 投稿",
                            snippet="恭喜，稿件已采用", verdict="过稿",
                            imap_uid="1001", received_at="2026-08-01 10:00:00"))
check("reply 插入", rid is not None and db.unread_count() == 1)
check("reply 去重（imap_uid+from_email）",
      db.insert_reply(Reply(from_email="ed1@x.com", imap_uid="1001", subject="重复")) is None)
db.mark_read(rid)
check("mark_read", db.unread_count() == 0 and len(db.list_replies()) == 1)

counts = db.counts()
check("counts 统计", counts["编辑总数"] == 2 and counts["文稿数"] == 1
      and counts["待回复"] == 0 and counts["过稿"] == 1 and counts["退稿"] == 0
      and counts["未读回信"] == 0)
sid2 = db.insert_submission(Submission(manuscript_id=mid, editor_id=eid2, status="待发"))
db.update_status(sid2, "已发")
act = db.recent_activity(10)
check("recent_activity 混合倒序", len(act) == 3 and act[0]["time"] >= act[-1]["time"]
      and act[-1]["kind"] == "回信"  # 回信时间最早，排最后
      and {a["kind"] for a in act} == {"投稿", "回信"})

db.delete_editor(eid2)
check("editor 删除", db.get_editor(eid2) is None)

# ---------- settings_store ----------
from app.settings_store import SettingsStore, provider_preset
from app.models import MailboxConfig, AuthorInfo

store = SettingsStore(db)
boxes = store.load_mailboxes()
check("load_mailboxes 补默认 6 个", len(boxes) == 6 and not boxes[0].enabled)
cfg = MailboxConfig(enabled=True, provider="QQ邮箱", address="a@qq.com", auth_code="x",
                    smtp_host="smtp.qq.com", smtp_port=465, smtp_ssl=True,
                    imap_host="imap.qq.com", imap_port=993, daily_limit=15)
store.save_mailbox(0, cfg)
loaded = store.load_mailboxes()[0]
check("save/load mailbox 往返", loaded.enabled and loaded.address == "a@qq.com"
      and loaded.daily_limit == 15 and loaded.smtp_ssl)
store.save_mailbox_count(8)
check("mailbox_count 往返", len(store.load_mailboxes()) == 8)
store.save_author(AuthorInfo(real_name="作者甲", pen_name="奶龙", phone="138", address="北京",
                             payment_info="支付宝"))
check("author 往返", store.load_author().pen_name == "奶龙")
store.save_strategy(True, 60, 25)
check("strategy 往返", store.get_strategy() == (True, 60, 25))
store.save_fetch_config(False, 15, 30)
check("fetch_config 往返", store.get_fetch_config() == (False, 15, 30))
_fresh_db = Database(db_path=os.path.join(tempfile.mkdtemp(prefix="nailong_fresh_"), "fresh.db"))
_fresh_store = SettingsStore(_fresh_db)
check("默认策略值", _fresh_store.get_strategy() == (True, 45, 30)
      and _fresh_store.get_fetch_config() == (True, 30, 45)
      and _fresh_store.get_theme() == "蔷薇粉")
store.set_theme("海岸蓝")
check("theme 往返", store.get_theme() == "海岸蓝")
check("provider_preset", provider_preset("QQ邮箱")[0] == "smtp.qq.com"
      and provider_preset("Outlook")[2] is False)

# ---------- classifier ----------
from app.classifier import classify_reply
check("classify 过稿", classify_reply("恭喜，您的稿件已被采用") == "过稿")
check("classify 退稿", classify_reply("很遗憾，本次未通过审核") == "退稿")
check("classify 需修改优先", classify_reply("稿件不错拟采用，但请先修改润色") == "需修改")
check("classify 其他", classify_reply("您好，请问有什么事？") == "其他")

# ---------- letter ----------
from app.letter import build_letter
subject, body = build_letter("测试小说", 12345, "短篇", "张三")
check("letter subject", subject == "投稿《测试小说》12345字 短篇")
check("letter body 称呼+祝颂语（无落款行）", body.startswith("尊敬的张三编辑")
      and "祝工作顺利，万事顺意！" in body
      and "笔名" not in body and "电话" not in body and "地址" not in body)

# ---------- docx_reader ----------
import docx
from app.docx_reader import read_docx_text, read_txt, count_cjk_words

docx_path = os.path.join(_tmp, "t.docx")
d = docx.Document()
d.add_paragraph("你好世界 hello world")
d.add_paragraph("第二段测试")
d.save(docx_path)
text = read_docx_text(docx_path)
check("read_docx_text", "你好世界" in text and "第二段测试" in text)
check("count_cjk_words", count_cjk_words("你好世界 hello world") == 6)  # 4 CJK + 2 词
txt_path = os.path.join(_tmp, "t.txt")
with open(txt_path, "wb") as f:
    f.write("中文文本".encode("gbk"))
check("read_txt gbk 回退", read_txt(txt_path) == "中文文本")

# ---------- theme ----------
from app.theme import THEMES, render_qss
check("5 套主题", len(THEMES) == 5 and THEMES["蔷薇粉"]["primary"] == "#D6336C")
qss = render_qss("墨紫")
check("render_qss 占位符替换", "#7048E8" in qss and "{primary}" not in qss)

# ---------- mailer/receiver 纯逻辑（不联网） ----------
from app.mailer import _display_name, _friendly_error, test_mailbox
check("display_name 留空用前缀", _display_name(MailboxConfig(address="hi@qq.com")) == "hi")
check("display_name 显式值", _display_name(MailboxConfig(address="hi@qq.com",
                                                       display_name="小龙")) == "小龙")
check("test_mailbox 缺参数", test_mailbox(MailboxConfig())[0] is False)
check("friendly_error 535", "授权码" in _friendly_error(Exception("535 Login Fail")))
from app.receiver import _strip_html, _extract_address
check("strip_html", _strip_html("<p>你好<b>编辑</b></p>").replace(" ", "") == "你好编辑")
check("extract_address 解码", _extract_address("=?utf-8?b?5byg5LiJ?= <ed@x.com>") == "ed@x.com")

# ---------- GUI：QApplication + MainWindow + navigate + 主题切换 ----------
from PySide6.QtWidgets import QApplication
from app.theme import apply_theme
from app.main_window import MainWindow

qapp = QApplication.instance() or QApplication([])
win = MainWindow(db, store)
check("窗口标题", win.windowTitle() == "奶龙投稿助手")
check("邮箱徽标", win.mail_badge.text() == "邮箱已配置 1/4")

for page_id in ("dashboard", "submit", "records", "replies",
                "manuscripts", "editors", "settings"):
    win.navigate(page_id)
    check(f"navigate {page_id}", win.stack.currentWidget() is win._pages[page_id])

for theme_name in THEMES:
    apply_theme(qapp, theme_name)
check("5 主题切换无异常", True)

settings_page = win._pages["settings"]
settings_page.refresh()
settings_page.save_all()
check("设置页保存+已保存提示", settings_page.save_hint.text() == "已保存")
win.data_changed.emit()
check("data_changed 信号刷新无异常", True)

print()
print(f"全部通过：{sum(PASS)}/{len(PASS)} 项")
